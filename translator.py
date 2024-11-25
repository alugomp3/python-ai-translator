import requests
import os
from docx import Document


subscription_key = 'YOUR_SUB_Key'
endpoint = 'https://Your_ENDpoint'
location = 'eastus2(Or Other)'
language_destination = 'pt-br'

path = '/translate'
constructd_url = endpoint + path
params = {
    'api-version': '3.0',
    'from': 'en',
    'to': [language_destination],

}

def translator_text(text, language_destination):
  path = '/translate'
  constructd_url = endpoint + path
  headers = {
      'Ocp-Apim-Subscription-key': subscription_key,
      'Ocp-Apim-Subscription-Region': location,
      'Content-type': 'application/json',
      'X-ClientTraceId': str(os.urandom(16))
      }
  body = {
      'text': text,
  }
  params = {
      'api-version': '3.0',
      'from': 'en',
      'to': [language_destination],
  }
  request = requests.post(constructd_url,params=params,headers=headers,json=body)
  response = request.json()
  return response[0]['translations'][0]['text']

def translator_document(text, language_destination):
  path = '/translate'
  constructd_url = endpoint + path
  headers = {
      'Ocp-Apim-Subscription-key': subscription_key,
      'Ocp-Apim-Subscription-Region': location,
      'Content-type': 'application/json',
      'X-ClientTraceId': str(os.random(16))
      }
  body = {
      'text': text,
  }
  params = {
      'api-version': '3.0',
      'from': 'en',
      'to': [language_destination],
  }
  request = requests.post(constructd_url,params=params,headers=headers,json=body)
  response = request.json()
  return response[0]['translations'][0]['text']

def translate_document(path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    translated_text = translator_text(paragraph.text, language_destination)
    full_text.append(translated_text)

  translated_doc = Document()
  for line in full_text:
    translated_doc.add_paragraph(line)

  path_translated = path.replace(".docx",f"{language_destination}.docx")
  translated_doc.save(path_translated)
  return path_translated

## Translate text call test
#translator_text('The blood that runs within my veins','pt-br')

#input_file = '/content/sample_data/anscombe.json'
## Upload Document to Drive
## Translate Document call test
#translate_document(input_file,'pt-br')
